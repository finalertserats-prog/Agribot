"""Shared docx styling + rendering engine for the AgriFriend doc pack.

Renders a simple block DSL (see content.py) into a styled .docx:
  ("h1", text) ("h2", text) ("h3", text)
  ("p", text)                      paragraph (supports **bold** inline)
  ("bul", [items])                 bullet list
  ("num", [items])                 numbered list
  ("table", headers, rows)         shaded-header table
  ("img", filename, caption)       centered image scaled to page width
  ("callout", kind, title, text)   note|warn|risk|ok shaded box
  ("kpi", [(value,label), ...])    stat row
  ("pagebreak",)

A cover page + static TOC (from h1/h2) + page-numbered footer are added
automatically. Static TOC avoids python-docx's live-field limitation.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DIAGRAMS = os.path.join(os.path.dirname(__file__), "..", "assets", "diagrams")

INK = RGBColor(0x26, 0x32, 0x38)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
GREEN_BG = "E8F5E9"
AMBER = RGBColor(0xEF, 0x6C, 0x00)
AMBER_BG = "FFF3E0"
RED = RGBColor(0xC6, 0x28, 0x28)
RED_BG = "FFEBEE"
BLUE = RGBColor(0x15, 0x65, 0xC0)
BLUE_BG = "E3F2FD"
GRAY = RGBColor(0x54, 0x6E, 0x7A)
GRAY_BG = "ECEFF1"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

CALLOUT = {
    "note": (BLUE, BLUE_BG, "NOTE"),
    "warn": (AMBER, AMBER_BG, "CAUTION"),
    "risk": (RED, RED_BG, "RISK"),
    "ok": (GREEN, GREEN_BG, "STRENGTH"),
    "honesty": (GRAY, GRAY_BG, "HONEST ASSESSMENT"),
}


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.append(borders)


def _set_cell_margins(table, top=60, bottom=60, left=120, right=120):
    tblPr = table._tbl.tblPr
    m = OxmlElement("w:tblCellMar")
    for edge, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tblPr.append(m)


def _runs_from_md(paragraph, text):
    """Very small **bold** inline parser."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        r = paragraph.add_run(part)
        r.font.color.rgb = INK
        if i % 2 == 1:
            r.bold = True


class DocBuilder:
    def __init__(self, meta):
        self.meta = meta
        self.doc = Document()
        self._base_style()
        self.toc_entries = []  # (level, text)
        self._blocks = []

    def _base_style(self):
        st = self.doc.styles["Normal"]
        st.font.name = "Calibri"
        st.font.size = Pt(10.5)
        st.font.color.rgb = INK
        for section in self.doc.sections:
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)
            section.top_margin = Inches(0.85)
            section.bottom_margin = Inches(0.85)

    # ---- footer with page number ----
    def _footer(self):
        section = self.doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{self.meta['doc_title']}   ·   ")
        run.font.size = Pt(8); run.font.color.rgb = GRAY
        fld1 = OxmlElement("w:fldSimple"); fld1.set(qn("w:instr"), "PAGE")
        run2 = p.add_run(); run2._r.append(fld1)
        run2.font.size = Pt(8); run2.font.color.rgb = GRAY
        run3 = p.add_run(" / "); run3.font.size = Pt(8); run3.font.color.rgb = GRAY
        fld2 = OxmlElement("w:fldSimple"); fld2.set(qn("w:instr"), "NUMPAGES")
        run4 = p.add_run(); run4._r.append(fld2)
        run4.font.size = Pt(8); run4.font.color.rgb = GRAY

    # ---- cover ----
    def _cover(self):
        m = self.meta
        band = self.doc.add_table(rows=1, cols=1)
        band.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = band.rows[0].cells[0]
        _shade(c, "1B5E20")
        _no_borders(band); _set_cell_margins(band, 400, 400, 300, 300)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(m["pack_name"]); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = WHITE
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(m["doc_title"]); r2.bold = True; r2.font.size = Pt(24); r2.font.color.rgb = WHITE
        p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(m["subtitle"]); r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0xC8, 0xE6, 0xC9)
        p4 = c.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run(m["variant"].upper() + " VERSION"); r4.bold = True
        r4.font.size = Pt(11); r4.font.color.rgb = WHITE

        self.doc.add_paragraph()
        # meta table
        t = self.doc.add_table(rows=0, cols=2)
        _no_borders(t); _set_cell_margins(t)
        info = [
            ("Document", m["doc_title"]),
            ("Part of", m["pack_name"] + f"  ({m['doc_index']} of 5)"),
            ("Version", m["variant"] + "  ·  " + m["date"]),
            ("Source (original)", "github.com/Shivaganesh-dev/agrifriend-bot"),
            ("Source (hardened)", "github.com/finalertserats-prog/Agribot  @ " + m["commit"]),
            ("Basis", f"48 unit tests · ~38% coverage · verdict: {m['verdict']}"),
        ]
        for k, v in info:
            row = t.add_row().cells
            _shade(row[0], GRAY_BG)
            rp = row[0].paragraphs[0]; rr = rp.add_run(k)
            rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = GRAY
            vp = row[1].paragraphs[0]; vr = vp.add_run(v)
            vr.font.size = Pt(9.5); vr.font.color.rgb = INK
        row = t.rows[0]
        # confidentiality
        self.doc.add_paragraph()
        cp = self.doc.add_paragraph()
        cr = cp.add_run("Prepared for internal review and stakeholder sharing. "
                        "Contains forward-looking and illustrative business estimates that are not guarantees.")
        cr.italic = True; cr.font.size = Pt(8.5); cr.font.color.rgb = GRAY
        self.doc.add_page_break()

    # ---- static TOC ----
    def _render_toc(self):
        h = self.doc.add_paragraph()
        hr = h.add_run("Contents"); hr.bold = True; hr.font.size = Pt(16); hr.font.color.rgb = GREEN
        for level, text in self.toc_entries:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3 if level == 2 else 0)
            r = p.add_run(("•  " if level == 2 else "") + text)
            r.font.size = Pt(11 if level == 1 else 10)
            r.bold = (level == 1)
            r.font.color.rgb = INK if level == 1 else GRAY
        self.doc.add_page_break()

    # ---- block renderers ----
    def _h(self, text, level):
        self.toc_entries.append((level, text))
        p = self.doc.add_paragraph()
        p.space_before = Pt(10)
        r = p.add_run(text); r.bold = True
        if level == 1:
            r.font.size = Pt(17); r.font.color.rgb = GREEN
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), "A5D6A7")
            pbdr.append(bottom); pPr.append(pbdr)
        elif level == 2:
            r.font.size = Pt(13); r.font.color.rgb = INK
        else:
            r.font.size = Pt(11); r.font.color.rgb = GRAY

    def _p(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _runs_from_md(p, text)

    def _list(self, items, numbered=False):
        style = "List Number" if numbered else "List Bullet"
        for it in items:
            p = self.doc.add_paragraph(style=style)
            _runs_from_md(p, it)

    def _table(self, headers, rows):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_cell_margins(t)
        # header
        for i, htext in enumerate(headers):
            c = t.rows[0].cells[i]; _shade(c, "2E7D32")
            p = c.paragraphs[0]; r = p.add_run(htext)
            r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for i, val in enumerate(row):
                if i < len(cells):
                    _shade(cells[i], "FFFFFF" if ri % 2 == 0 else "F5F7F5")
                    p = cells[i].paragraphs[0]
                    _runs_from_md(p, str(val))
                    for run in p.runs:
                        run.font.size = Pt(9)
        t.style = "Table Grid"
        self.doc.add_paragraph()

    def _img(self, filename, caption):
        path = os.path.join(DIAGRAMS, filename)
        if not os.path.exists(path):
            self._p(f"[missing diagram: {filename}]")
            return
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(6.4))
        if caption:
            cp = self.doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run("Figure. " + caption)
            cr.italic = True; cr.font.size = Pt(8.5); cr.font.color.rgb = GRAY
        self.doc.add_paragraph()

    def _callout(self, kind, title, text):
        color, bg, label = CALLOUT[kind]
        t = self.doc.add_table(rows=1, cols=1); _no_borders(t)
        _set_cell_margins(t, 120, 120, 160, 160)
        c = t.rows[0].cells[0]; _shade(c, bg)
        p = c.paragraphs[0]
        lr = p.add_run(f"{label}" + (f" — {title}" if title else ""))
        lr.bold = True; lr.font.size = Pt(9.5); lr.font.color.rgb = color
        bp = c.add_paragraph(); _runs_from_md(bp, text)
        for r in bp.runs:
            r.font.size = Pt(9.5)
        # left accent border
        tcPr = c._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "0"); left.set(qn("w:color"),
                                               "%02X%02X%02X" % (color[0], color[1], color[2]))
        borders.append(left); tcPr.append(borders)
        self.doc.add_paragraph()

    def _kpi(self, items):
        t = self.doc.add_table(rows=1, cols=len(items)); _no_borders(t)
        _set_cell_margins(t, 100, 100, 100, 100)
        for i, (value, label) in enumerate(items):
            c = t.rows[0].cells[i]; _shade(c, GREEN_BG)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vr = p.add_run(str(value)); vr.bold = True; vr.font.size = Pt(20); vr.font.color.rgb = GREEN
            lp = c.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lr = lp.add_run(label); lr.font.size = Pt(8.5); lr.font.color.rgb = GRAY
        self.doc.add_paragraph()

    def render(self, blocks):
        self._cover()
        # first pass to collect TOC without emitting (we emit toc then body,
        # so gather headings first)
        for b in blocks:
            if b[0] in ("h1", "h2"):
                self.toc_entries.append((1 if b[0] == "h1" else 2, b[1]))
        self._render_toc()
        self.toc_entries = []  # reset (already rendered)
        for b in blocks:
            kind = b[0]
            if kind == "h1": self._h(b[1], 1)
            elif kind == "h2": self._h(b[1], 2)
            elif kind == "h3": self._h(b[1], 3)
            elif kind == "p": self._p(b[1])
            elif kind == "bul": self._list(b[1], numbered=False)
            elif kind == "num": self._list(b[1], numbered=True)
            elif kind == "table": self._table(b[1], b[2])
            elif kind == "img": self._img(b[1], b[2] if len(b) > 2 else "")
            elif kind == "callout": self._callout(b[1], b[2], b[3])
            elif kind == "kpi": self._kpi(b[1])
            elif kind == "pagebreak": self.doc.add_page_break()
        self._footer()
        return self.doc

    def save(self, path):
        self.doc.save(path)
        return path
